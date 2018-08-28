# ===================================
# Filename: __main__.py
# Purpose: To setup and start the docx-diary program.
#
#
# docx-diary
# Copyright (C) 2018  Joshua Peter Booth
#
# This file is part of docx-diary.
#
# docx-diary is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# docx-diary is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with docx-diary (see LICENSE.md).
# If not, see <http://www.gnu.org/licenses/>.
#
# Contact me:
# Email: joshb00th@icloud.com
# ===================================

from _logger import logger
import os
import datetime
import argparse
import docx
from docx.shared import Pt

logger.disabled = False


def main():
    """ Start of the program. """
    if not check_change_directory('diary'):
        print("Checking current folder as no diary folder was found...")
    parser_description = "View, create and change data from a text file."
    parser = argparse.ArgumentParser(description=parser_description)
    group = parser.add_mutually_exclusive_group()
    parser.add_argument("-d", "--date", action="store", nargs=3, type=int,
                        metavar=("DAY", "MONTH", "YEAR"),
                        help="View data from a date")
    group.add_argument("-n", "--new", action="store_true",
                       help="Create a new data entry")
    group.add_argument("-e", "--edit", action="store", nargs=3, type=int,
                       metavar=("DAY", "MONTH", "YEAR"),
                       help="Edit an existing data entry")
    args = parser.parse_args()

    # Create a new data entry for the current date and time
    if args.new:
        file_name = "{} {} {} {}.docx".format(get_current_date()["day_name"],
                                              get_current_date()["month_name"],
                                              get_current_date()["day"],
                                              get_current_date()["year"])
        if is_file(file_name):
            print("Do you want to overwrite your current message? (y/n)")
            overwrite = input(">> ")
            if overwrite.lower() == "y":
                new_data()
        else:
            new_data()
    # View data from a date
    elif args.date is not None:
        day, month, year = args.date
        date_format = {'day': day, 'month': month, 'year': year}
        view_data(date_format)
    # Edit a data entry
    elif args.edit is not None:
        day, month, year = args.edit
        date_format = {'day': day, 'month': month, 'year': year}
        edit_data(date_format)
    else:
        logger.error("No arguments were entered!")
        print("No arguments were entered!")
        print("Please enter -h or --help for help.")


def is_file(file):
    """
    Check if a file is an openable word document (.docx).

    :param str file: The file to check.
    :return bool True | False:
    """
    return True if get_data(file) is not None else False


def get_current_date():
    """
    Return the current date in a dict format.

    :returns dict:
             keys: (int) "year": The full year.
                   (int) "month": The month number.
                   (str) "month_name": The full month name.
                   (int) "day": The day number of the month.
                   (str) "day_name": The full name of the day.
                   (int) "hour": The hour of the day.
                   (int) "minute": The minute of the hour.
                   (str) "period": The period being either AM or PM.
    """
    date = datetime.datetime.now()

    return {
        "year": int(date.strftime("%Y")),
        "month": int(date.strftime("%m")),
        "month_name": date.strftime("%B"),
        "day": int(date.strftime("%d")),
        "day_name": date.strftime("%A"),
        "hour": int(date.strftime("%I")),
        "minute": int(date.strftime("%M")),
        "period": date.strftime("%p")
    }


def get_data(filename):
    """
    Return all the data from a file with the given filename.

    :param str filename: The name of the file in which to get data.
    :returns str full_data | None: Gets the data or else nothing.
    """
    try:
        doc = docx.Document(filename)
        full_data = []
        for para in doc.paragraphs:
            full_data.append(para.text)
        return full_data
    except Exception as e:
        logger.error(e)


def get_filename(date):
    """
    Get the filename from a given date.

    :param dict date: The data that will be used to get the filename.
    :return None:
    """
    for item in date:
        try:
            date[item] = int(date[item])
        except ValueError:
            print("Failed to convert to integer.")

    time = datetime.datetime(date['year'], date['month'], date['day'])
    file_name = str(time.date().strftime('%A %B %d %Y') + '.docx').split(' ')
    file_name[2] = str(int(file_name[2]))
    file_name = " ".join(file_name)
    return file_name


def view_data(date):
    """
    Displays the data from a given date.

    :param dict date:
    :return None:
    """
    file_name = get_filename(date)
    try:
        if check_file(file_name) is not False:
            data = get_data(file_name)
        with open(file_name, 'r'):
            day_of_week = get_date_data(file_name)['day_of_week']
            month = get_date_data(file_name)['month']
            day = get_date_data(file_name)['day']
            year = get_date_data(file_name)['year']

            file_date = "{} {} {} {}.docx".format(day_of_week, month, day,
                                                  year)
            if file_name == file_date:
                print("-"*45)
                print("Date: ", end=" ")
                for item in data:
                    print(item)
                    if item[len(item)-1] == "M":
                        print("Message: ", end="\n")
                print("-"*45)
            else:
                logger.error("File_date '{}' does not match the file_name '{}'"
                             .format(file_date, file_name))
                return
    except FileNotFoundError:
        logger.error("File '{}' could not be found!".format(file_name))
        print("There is no data for the date '{}/{}/{}'.".format(date['day'],
                                                                 date['month'],
                                                                 date['year']))


def new_data():
    """
    Create a new data entry for the current date and time.

    :return None:
    """
    current_date = list(value for value in get_current_date().values())
    year, month, month_name, day, day_name, hour, minute, period = current_date

    file_time = "{}:{}".format(hour, minute)
    file_name = "{} {} {} {}.docx".format(day_name, month_name, day, year)
    file_date = "{}, {} {}, {}".format(day_name, month_name, day, year)

    document = docx.Document()
    style = document.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    with open(file_name, 'w'):
        message = input("Date: {}\nEnter your message: ".format(file_date))

        date = document.add_paragraph()
        date.style = document.styles['Normal']
        data = document.add_paragraph()
        data.style = document.styles['Normal']
        date.add_run("{} {} {}".format(file_date, file_time, period))
        data.add_run(message)
        document.save(file_name)


def edit_data(date):
    """
    Opens the data file from a given date to be edited.

    :param date: (dict)
    :return: (None)
    """
    try:
        file_name = get_filename(date)
        os.startfile("{}\{}".format(os.getcwd(), file_name))
    except FileNotFoundError:
        logger.error("File could not be found: {}".format(get_filename(date)))
        date = "{}/{}/{}".format(date['day'], date['month'], date['year'])
        print("There is no data you can edit for the date '{}'.".format(date))


def check_change_directory(directory):
    """
    Checks if the directory can be changed to by trying to change the directory
    to the given directory name.

    :param str directory: The directory to check is valid.
    :returns bool True | None:
    """
    try:
        os.chdir(directory)
        return True
    except FileNotFoundError:
        logger.error("Directory '{}' not found.".format(directory))
        return False


def check_file(data_file):
    """
    Checks a given file to see if it is in a valid format by having at least 2
    lines in the file.

    :param str data_file: The data file to check is valid.
    :return bool True | False:
    """
    if is_file(data_file):
        with open(data_file, 'r'):
            doc = docx.Document(data_file)
            counter = 0
            for para in doc.paragraphs:
                if para.text == "":
                    break
                counter += 1
            if counter == 1:
                logger.error("No message to read from.")
            elif counter < 1:
                logger.error("No date or message to read from.")
            else:
                return True
    else:
        return False


def get_date_data(file_name):
    """
    Returns a dictionary of date information in the given file.

    :param str file_name: The file to get the date from.
    :returns dict or None: (dict) or (None)
             keys: "day_of_week" (str)
                   "month" (str)
                   "day" (str)
                   "year" (str)
                   "time" (str)
                   "period" (str)
    """
    try:
        with open(file_name, 'r'):
            date = get_data(file_name)[0].split(' ')
            return {
                'day_of_week': date[0].strip(','),
                'month': date[1].strip(','),
                'day': date[2].strip(','),
                'year': date[3],
                'time': date[4],
                'period': date[5][0:2]
            }
    except FileNotFoundError:
        logger.error("File '{}' not found.".format(file_name))


if __name__ == '__main__':
    main()
